"""
DocMind AI v2.2 — Stable / Graceful Degradation Release
"""

import streamlit as st
import hashlib
import html
import io
import json
import logging
import os
import re
from datetime import datetime
from pathlib import Path

from docmind_core import (
    DEFAULT_BASE_URL,
    DEFAULT_CHUNK_OVERLAP,
    DEFAULT_CHUNK_SIZE,
    MAX_CHUNKS,
    MAX_TRANSFORM_CONTINUATIONS,
    TRANSFORM_CHUNK_SIZE,
    AIConfigurationError,
    AIResponseError,
    DocumentReadError,
    SearchOutcome,
    basic_text_diff,
    chunk_scope_notice,
    complete_with_continuations,
    estimate_token_usage,
    extract_completion_result,
    extract_token_usage,
    log_failure,
    make_history_entry,
    parse_document,
    resolve_ai_config,
    run_optional_ai,
    safe_error_message,
    search_web,
    split_text_into_chunks,
)


LOG_LEVEL = getattr(
    logging, os.environ.get("DOCMIND_LOG_LEVEL", "INFO").upper(), logging.INFO
)
logging.basicConfig(
    level=LOG_LEVEL,
    format="%(asctime)s %(levelname)s %(name)s: %(message)s",
)
logger = logging.getLogger("docmind.app")
MAX_BASIC_WORD_EXPORT_CHARS = 500_000
MAX_SESSION_PARSE_CACHE_BYTES = 80 * 1024 * 1024

# ─────────────────────────────────────────────
#  页面配置
# ─────────────────────────────────────────────
st.set_page_config(
    page_title="DocMind AI · 智能文档助手",
    page_icon="🧠",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ─────────────────────────────────────────────
#  全局CSS — 电脑保持原样，手机自动适配
# ─────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap');

html, body, [class*="css"], p, div, span, label {
    font-family: 'Inter', sans-serif !important;
}
#MainMenu, footer, header { visibility: hidden; }
.stApp { background: #f8f9fc !important; }

/* ── 电脑端布局（默认） ── */
.main .block-container {
    padding-top: 1.5rem !important;
    padding-bottom: 2rem !important;
    max-width: 1100px !important;
}

/* ══════════════════════════════════════════
   手机端适配（宽度 ≤ 768px 自动触发）
   电脑端完全不受影响
   ══════════════════════════════════════════ */
@media (max-width: 768px) {
    .main .block-container {
        padding: 0.5rem 0.4rem 2rem 0.4rem !important;
        max-width: 100% !important;
    }
    /* 品牌栏手机版缩小 */
    .brand-title { font-size: 1.1rem !important; }
    .brand-sub   { display: none !important; }
    .brand-bar   { padding: 12px 16px !important; border-radius: 10px !important; }
    /* 操作按钮手机版2列 */
    div[data-testid="column"] { min-width: 45% !important; }
    /* 侧边栏提示 */
    .mobile-key-hint { display: block !important; }
    .desktop-only    { display: none !important; }
    /* Tab文字缩小 */
    .stTabs [data-baseweb="tab"] {
        font-size: 0.75rem !important;
        padding: 6px 8px !important;
    }
    /* 文件上传区 */
    [data-testid="stFileUploader"] { padding: 4px !important; }
}

/* 手机提示条默认隐藏，只在手机显示 */
.mobile-key-hint { display: none; }

/* ── 品牌栏 ── */
.brand-bar {
    background: linear-gradient(135deg, #1a1f36 0%, #2d3561 100%);
    padding: 22px 32px;
    border-radius: 16px;
    margin-bottom: 20px;
    box-shadow: 0 4px 24px rgba(26,31,54,0.2);
}
.brand-inner { display: flex; align-items: center; gap: 14px; }
.brand-logo  { font-size: 2.2rem; line-height: 1; }
.brand-title {
    font-size: 1.55rem; font-weight: 700;
    color: #ffffff !important; letter-spacing: -0.5px; margin: 0;
}
.brand-sub {
    font-size: 0.82rem; color: rgba(255,255,255,0.5) !important; margin: 3px 0 0 0;
}
.brand-badge {
    margin-left: auto;
    background: rgba(99,179,237,0.18); border: 1px solid rgba(99,179,237,0.35);
    color: #63b3ed !important; font-size: 0.72rem; font-weight: 700;
    padding: 4px 14px; border-radius: 20px; letter-spacing: 1px;
}

/* ── 侧边栏 ── */
[data-testid="stSidebar"] > div:first-child { background: #1a1f36 !important; }
[data-testid="stSidebar"] { background: #1a1f36 !important; }
[data-testid="stSidebar"] p,
[data-testid="stSidebar"] span,
[data-testid="stSidebar"] label,
[data-testid="stSidebar"] div,
[data-testid="stSidebar"] h1,
[data-testid="stSidebar"] h2,
[data-testid="stSidebar"] h3 { color: rgba(255,255,255,0.88) !important; }
[data-testid="stSidebar"] .stTextInput input,
[data-testid="stSidebar"] .stSelectbox select {
    background: rgba(255,255,255,0.08) !important;
    color: white !important; border-color: rgba(255,255,255,0.15) !important;
}
[data-testid="stSidebar"] hr { border-color: rgba(255,255,255,0.12) !important; }
[data-testid="stSidebar"] .stButton > button {
    background: rgba(255,255,255,0.1) !important;
    color: white !important; border: 1px solid rgba(255,255,255,0.2) !important;
}

/* ── 按钮 ── */
.stButton > button {
    border-radius: 10px !important; font-weight: 600 !important;
    font-size: 0.88rem !important; padding: 0.45rem 1rem !important;
    transition: all 0.18s !important;
}
.stButton > button[kind="primary"] {
    background: #4f6ef7 !important; border: none !important; color: white !important;
}
.stButton > button[kind="primary"]:hover {
    background: #3d5ce0 !important; transform: translateY(-1px) !important;
    box-shadow: 0 4px 14px rgba(79,110,247,0.35) !important;
}
.stDownloadButton > button { border-radius: 10px !important; font-weight: 600 !important; }

/* ── Tab样式 ── */
.stTabs [data-baseweb="tab-list"] {
    gap: 4px; background: transparent; border-bottom: 2px solid #e8eaf0;
}
.stTabs [data-baseweb="tab"] {
    border-radius: 8px 8px 0 0 !important; font-weight: 500 !important;
    font-size: 0.9rem !important; padding: 8px 18px !important; color: #8892a4 !important;
}
.stTabs [aria-selected="true"] {
    color: #4f6ef7 !important; background: #f0f3ff !important;
    border-bottom: 2px solid #4f6ef7 !important;
}

/* ── 文件上传 ── */
[data-testid="stFileUploader"] {
    border: 2px dashed #c8cfe0 !important;
    border-radius: 12px !important; background: #f8f9fc !important; padding: 8px !important;
}
[data-testid="stFileUploader"]:hover { border-color: #4f6ef7 !important; }

/* ── 卡片/信息框 ── */
.info-card {
    background: #f0f3ff; border: 1px solid #d0d8ff;
    border-left: 4px solid #4f6ef7; border-radius: 10px;
    padding: 12px 16px; font-size: 0.85rem; color: #2d3748;
    margin: 8px 0 12px 0; line-height: 1.6;
}
.warn-card {
    background: #fffbeb; border: 1px solid #fcd34d;
    border-left: 4px solid #f59e0b; border-radius: 10px;
    padding: 12px 16px; font-size: 0.85rem; color: #78350f;
    margin: 8px 0 12px 0; line-height: 1.6;
}

/* ── 文件chip ── */
.file-chip {
    display: inline-flex; align-items: center; gap: 6px;
    background: #f0f3ff; border: 1px solid #d0d8ff;
    border-radius: 8px; padding: 6px 12px; margin: 3px 0;
    font-size: 0.84rem; color: #1a1f36;
}

/* ── 搜索结果 ── */
.search-item {
    border-left: 3px solid #4f6ef7; padding: 10px 14px;
    margin-bottom: 10px; background: #f8f9fc;
    border-radius: 0 8px 8px 0; border: 1px solid #e8eaf0;
    border-left: 3px solid #4f6ef7;
}
.search-title { font-weight: 600; font-size: 0.9rem; color: #1a1f36; }
.search-url   { font-size: 0.75rem; color: #8892a4; margin: 3px 0; }
.search-snip  { font-size: 0.85rem; color: #4a5568; line-height: 1.55; }

/* ── 标签 ── */
.tag {
    display: inline-block; background: #eef0ff; color: #4f6ef7;
    font-size: 0.73rem; font-weight: 600; padding: 3px 10px;
    border-radius: 20px; margin-right: 6px;
}

/* ── 统计 ── */
.stat-box { text-align: center; padding: 12px 8px; }
.stat-number { font-size: 1.8rem; font-weight: 700; color: white !important; line-height: 1; }
.stat-label  { font-size: 0.75rem; color: rgba(255,255,255,0.5) !important; margin-top: 4px; font-weight: 500; }

/* ── 进度条 ── */
.stProgress > div > div { background: #4f6ef7 !important; border-radius: 4px !important; }

/* ── 登录页 ── */
.login-card {
    background: #fff; border: 1px solid #e2e8f0; border-radius: 20px;
    padding: 40px 36px; box-shadow: 0 8px 40px rgba(79,110,247,0.12);
    max-width: 400px; margin: 60px auto 0;
}

/* ── 对比结果 ── */
.diff-section {
    background: #fff; border: 1px solid #e2e8f0; border-radius: 12px;
    padding: 16px 20px; margin: 8px 0;
    box-shadow: 0 1px 4px rgba(0,0,0,0.04);
}
</style>
""", unsafe_allow_html=True)

# ─────────────────────────────────────────────
#  服务端配置与 Session state
# ─────────────────────────────────────────────
def _load_server_secrets():
    try:
        return st.secrets.to_dict()
    except Exception as exc:
        if type(exc).__name__ == "StreamlitSecretNotFoundError":
            logger.info("Streamlit secrets file is not configured; using environment variables")
        else:
            logger.warning("Unable to load Streamlit secrets [%s]", type(exc).__name__)
        return {}


SERVER_SECRETS = _load_server_secrets()


def _server_value(*names, default=""):
    for name in names:
        value = SERVER_SECRETS.get(name)
        if value not in (None, ""):
            return value
        value = os.environ.get(name)
        if value not in (None, ""):
            return value
    return default


SERVER_DEEPSEEK_KEY = str(_server_value("DEEPSEEK_API_KEY", default=""))
SERVER_OPENAI_KEY = str(_server_value("OPENAI_API_KEY", default=""))
if SERVER_DEEPSEEK_KEY:
    SERVER_API_KEY = SERVER_DEEPSEEK_KEY
    SERVER_BASE_URL = str(_server_value("DEEPSEEK_BASE_URL", default=DEFAULT_BASE_URL))
    SERVER_MODEL = str(_server_value("DEEPSEEK_MODEL", default="deepseek-chat"))
elif SERVER_OPENAI_KEY:
    SERVER_API_KEY = SERVER_OPENAI_KEY
    SERVER_BASE_URL = str(
        _server_value("OPENAI_BASE_URL", default="https://api.openai.com/v1")
    )
    SERVER_MODEL = str(_server_value("OPENAI_MODEL", default="gpt-4o-mini"))
else:
    SERVER_API_KEY = ""
    SERVER_BASE_URL = DEFAULT_BASE_URL
    SERVER_MODEL = "deepseek-chat"
SERVER_TAVILY_KEY = str(_server_value("TAVILY_API_KEY", "TAVILY_KEY", default=""))
SERVER_BOCHA_KEY = str(_server_value("BOCHA_API_KEY", "BOCHA_KEY", default=""))


def _init_state():
    # v2.0 曾把服务端密钥写入 Session；热重载时主动移除这些旧字段。
    for obsolete_key in ("api_key", "base_url", "tavily_key", "bocha_key"):
        st.session_state.pop(obsolete_key, None)
    defaults = {
        "byok_api_key":  "",
        "byok_provider": "DeepSeek",
        "byok_base_url": DEFAULT_BASE_URL,
        "model":         SERVER_MODEL,
        "total_calls":   0,
        "total_tokens":  0,
        "prompt_tokens": 0,
        "completion_tokens": 0,
        "estimated_calls": 0,
        "actual_usage_calls": 0,
        "history":       [],
        "web_enabled":   False,
        "byok_bocha_key": "",
        "byok_tavily_key": "",
        "network_mode":  "auto",
        "logged_in":     False,
        "username":      "",
        "chat_messages": [],
        "selected_op":   "📝 智能摘要",
        "compare_state": None,
        "parsed_cache": {},
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

_init_state()

# ─────────────────────────────────────────────
#  登录系统（可选，设置 DOCMIND_USERS 才启用）
# ─────────────────────────────────────────────
USERS_CONFIG = _server_value("DOCMIND_USERS", default="")
LOGIN_CONFIG_ERROR = False
try:
    if isinstance(USERS_CONFIG, dict):
        USERS = {str(key): str(value) for key, value in USERS_CONFIG.items()}
    elif USERS_CONFIG:
        parsed_users = json.loads(str(USERS_CONFIG))
        if not isinstance(parsed_users, dict):
            raise TypeError("DOCMIND_USERS must be a JSON object")
        USERS = {str(key): str(value) for key, value in parsed_users.items()}
    else:
        USERS = {}
except (json.JSONDecodeError, TypeError, ValueError) as exc:
    logger.error("Invalid DOCMIND_USERS configuration [%s]", type(exc).__name__)
    USERS = {}
    LOGIN_CONFIG_ERROR = True
LOGIN_ON = bool(USERS)

def render_login():
    st.markdown("""
    <div class='login-card'>
      <div style='text-align:center;margin-bottom:24px'>
        <div style='font-size:3rem'>🧠</div>
        <h2 style='color:#1e293b;margin:8px 0 4px;font-size:1.5rem;font-weight:700'>DocMind AI</h2>
        <p style='color:#94a3b8;font-size:0.85rem;margin:0'>智能文档处理平台</p>
      </div>
    </div>
    """, unsafe_allow_html=True)
    col = st.columns([1, 2, 1])[1]
    with col:
        st.text_input("", placeholder="用户名", key="login_user", label_visibility="collapsed")
        st.text_input("", placeholder="密码", type="password", key="login_pass", label_visibility="collapsed")
        if st.button("登 录", width="stretch", type="primary"):
            u = st.session_state.login_user
            p = st.session_state.login_pass
            if u in USERS and USERS[u] == p:
                st.session_state.logged_in = True
                st.session_state.username  = u
                st.rerun()
            else:
                st.error("用户名或密码错误")

# ─────────────────────────────────────────────
#  文件读取
# ─────────────────────────────────────────────
def read_uploaded_document(file):
    data = file.getvalue()
    cache_key = hashlib.sha256(file.name.encode("utf-8") + b"\0" + data).hexdigest()
    cached = st.session_state.parsed_cache.get(cache_key)
    if cached is not None:
        return cached
    parsed = parse_document(file.name, data)
    parsed_size = parsed.size_bytes + len(parsed.text.encode("utf-8"))
    cached_size = sum(
        item.size_bytes + len(item.text.encode("utf-8"))
        for item in st.session_state.parsed_cache.values()
    )
    while st.session_state.parsed_cache and (
        len(st.session_state.parsed_cache) >= 8
        or cached_size + parsed_size > MAX_SESSION_PARSE_CACHE_BYTES
    ):
        oldest_key = next(iter(st.session_state.parsed_cache))
        removed = st.session_state.parsed_cache.pop(oldest_key)
        cached_size -= removed.size_bytes + len(removed.text.encode("utf-8"))
    if parsed_size <= MAX_SESSION_PARSE_CACHE_BYTES:
        st.session_state.parsed_cache[cache_key] = parsed
    return parsed


def read_uploaded(file) -> str:
    return read_uploaded_document(file).text

# ─────────────────────────────────────────────
#  AI调用
# ─────────────────────────────────────────────
def get_active_ai_config():
    return resolve_ai_config(
        server_api_key=SERVER_API_KEY,
        server_base_url=SERVER_BASE_URL,
        byok_api_key=st.session_state.byok_api_key,
        byok_base_url=st.session_state.byok_base_url,
        model=st.session_state.model,
        server_model=SERVER_MODEL,
    )


def has_ai_config() -> bool:
    try:
        get_active_ai_config()
        return True
    except AIConfigurationError:
        return False


def call_ai(messages: list, stream=True, max_tokens=3000):
    from openai import OpenAI
    import httpx

    config = get_active_ai_config()
    client = OpenAI(
        api_key=config.api_key,
        base_url=config.base_url,
        timeout  = httpx.Timeout(connect=10, read=120, write=30, pool=10),
    )
    return client.chat.completions.create(
        model=config.model,
        messages=messages,
        max_tokens=max_tokens,
        stream=stream,
    )


def _record_ai_usage(messages, output, usage=None):
    resolved_usage = usage or estimate_token_usage(messages, output)
    st.session_state.total_calls += 1
    st.session_state.prompt_tokens += resolved_usage.prompt_tokens
    st.session_state.completion_tokens += resolved_usage.completion_tokens
    st.session_state.total_tokens += resolved_usage.total_tokens
    if resolved_usage.estimated:
        st.session_state.estimated_calls += 1
    else:
        st.session_state.actual_usage_calls += 1


def complete_reply_result(messages, max_tokens=3000, track_usage=True):
    response = call_ai(messages, stream=False, max_tokens=max_tokens)
    result = extract_completion_result(response)
    if track_usage:
        _record_ai_usage(messages, result.text, extract_token_usage(response))
    return result


def complete_reply(messages, max_tokens=3000, track_usage=True) -> str:
    return complete_reply_result(messages, max_tokens, track_usage).text


def complete_with_continuation(messages, max_tokens=3500):
    """检测 length 截断并有限续写；返回文本和是否仍不完整。"""
    outcome = complete_with_continuations(
        lambda current_messages: complete_reply_result(
            current_messages,
            max_tokens=max_tokens,
        ),
        messages,
        max_continuations=MAX_TRANSFORM_CONTINUATIONS,
    )
    return outcome.text, outcome.incomplete


def stream_reply(messages, placeholder, max_tokens=3000) -> str:
    full = ""
    usage = None
    finish_reason = ""
    for chunk in call_ai(messages, stream=True, max_tokens=max_tokens):
        delta = ""
        if getattr(chunk, "choices", None):
            choice = chunk.choices[0]
            delta = choice.delta.content or ""
            finish_reason = getattr(choice, "finish_reason", None) or finish_reason
        full += delta
        chunk_usage = extract_token_usage(chunk)
        if chunk_usage is not None:
            usage = chunk_usage
        placeholder.markdown(clean_md(full) + "▌")
    if not full.strip():
        raise AIResponseError("AI 服务返回了空内容，请稍后重试或更换模型。")
    placeholder.markdown(clean_md(full))
    _record_ai_usage(messages, full, usage)
    if finish_reason == "length":
        st.warning("AI 输出达到长度上限，结果可能不完整。请缩小处理范围后重试。")
    return full

# ─────────────────────────────────────────────
#  联网搜索
# ─────────────────────────────────────────────
def web_search(query: str) -> SearchOutcome:
    return search_web(
        query=query,
        network_mode=st.session_state.network_mode,
        tavily_key=SERVER_TAVILY_KEY or st.session_state.byok_tavily_key,
        bocha_key=SERVER_BOCHA_KEY or st.session_state.byok_bocha_key,
    )


def render_search_feedback(outcome: SearchOutcome):
    for notice in outcome.notices:
        st.caption(f"ℹ️ {notice}")
    if outcome.error:
        st.warning(outcome.error)

# ─────────────────────────────────────────────
#  工具函数
# ─────────────────────────────────────────────
def clean_md(text: str) -> str:
    text = re.sub(r'```[^\n]*\n([\s\S]*?)```', r'\1', text)
    text = re.sub(r'#{1,6}\s*', '', text)
    text = re.sub(r'\*{1,3}(.+?)\*{1,3}', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    text = re.sub(r'^[-*+]\s+', '• ', text, flags=re.MULTILINE)
    text = re.sub(r'(?<!\w)[*#]+(?!\w)', '', text)
    return text.strip()

def to_docx_bytes(content: str, title: str = "AI处理结果") -> bytes:
    import docx as _docx
    from docx.shared import RGBColor
    doc = _docx.Document()
    h   = doc.add_heading(title, level=1)
    h.runs[0].font.color.rgb = RGBColor(0x4f, 0x6e, 0xf7)
    doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph("─" * 40)
    doc.add_paragraph(clean_md(content))
    buf = io.BytesIO(); doc.save(buf)
    return buf.getvalue()

def save_history(op, filename, result):
    st.session_state.history.insert(
        0,
        make_history_entry(op, filename, result, datetime.now().strftime("%H:%M")),
    )
    if len(st.session_state.history) > 50:
        st.session_state.history = st.session_state.history[:50]


DOCUMENT_SYSTEM_PROMPT = (
    "你是专业的文档分析助手，擅长处理商业文件、合同、报告和数据表格。"
    "输出内容清晰、结构化、专业；不要编造原文未出现的信息。"
)


def process_long_document(
    content,
    file_name,
    instruction,
    operation_mode,
    placeholder,
    search_context="",
    status_callback=None,
):
    """短文档单次处理；长文档按操作类型执行 Map-Reduce 或分块拼接。"""
    overlap = 0 if operation_mode == "concatenate" else DEFAULT_CHUNK_OVERLAP
    chunk_size = TRANSFORM_CHUNK_SIZE if operation_mode == "concatenate" else DEFAULT_CHUNK_SIZE
    chunking = split_text_into_chunks(
        content,
        chunk_size=chunk_size,
        overlap=overlap,
        max_chunks=MAX_CHUNKS,
    )
    if not chunking.chunks:
        raise DocumentReadError(f"{file_name} 没有可处理的文本内容。")

    if len(chunking.chunks) == 1 and not chunking.truncated:
        messages = [
            {"role": "system", "content": DOCUMENT_SYSTEM_PROMPT},
            {
                "role": "user",
                "content": (
                    f"{instruction}\n\n---文件：{file_name}---\n"
                    f"{chunking.chunks[0]}{search_context}"
                ),
            },
        ]
        if operation_mode == "concatenate":
            output, incomplete = complete_with_continuation(messages)
            placeholder.markdown(clean_md(output))
            return output, chunking, incomplete
        return stream_reply(messages, placeholder), chunking, False

    local_results = []
    incomplete_output = False
    total_chunks = len(chunking.chunks)
    for index, chunk in enumerate(chunking.chunks, start=1):
        if status_callback:
            status_callback(index, total_chunks)
        if operation_mode == "concatenate":
            map_instruction = (
                f"{instruction}\n\n这是文件的第 {index}/{total_chunks} 块。"
                "仅处理当前块，保持内容顺序，不省略，不添加跨块总结。"
            )
            max_tokens = 3800
        else:
            map_instruction = (
                f"{instruction}\n\n这是文件的第 {index}/{total_chunks} 块。"
                "请只基于当前块给出简洁、结构化的局部结果，保留关键名称、数字、日期、金额和风险。"
            )
            max_tokens = 1400
        messages = [
            {"role": "system", "content": DOCUMENT_SYSTEM_PROMPT},
            {
                "role": "user",
                "content": f"{map_instruction}\n\n---当前文本块---\n{chunk}",
            },
        ]
        if operation_mode == "concatenate":
            local_result, still_incomplete = complete_with_continuation(
                messages,
                max_tokens=max_tokens,
            )
            incomplete_output = incomplete_output or still_incomplete
            local_results.append(local_result)
        else:
            local_results.append(complete_reply(messages, max_tokens=max_tokens))

    if operation_mode == "concatenate":
        final_text = "\n\n".join(local_results)
        placeholder.markdown(clean_md(final_text))
        return final_text, chunking, incomplete_output

    mapped_text = "\n\n".join(
        f"[文本块 {index}]\n{result}"
        for index, result in enumerate(local_results, start=1)
    )
    reduce_messages = [
        {"role": "system", "content": DOCUMENT_SYSTEM_PROMPT},
        {
            "role": "user",
            "content": (
                f"原任务：{instruction}\n\n"
                "以下是同一文件各文本块的局部结果。请综合、去重并形成一份完整最终结果；"
                "保留关键名称、数字、日期、金额、例外条件与风险，不要遗漏块间矛盾。"
                f"{search_context}\n\n---局部结果集合---\n{mapped_text}"
            ),
        },
    ]
    return stream_reply(reduce_messages, placeholder), chunking, False


def prepare_comparison_content(file_name, content, status_callback=None):
    """为长文档生成结构化分块表示，供最终双文档对比。"""
    chunking = split_text_into_chunks(content)
    if not chunking.chunks:
        raise DocumentReadError(f"{file_name} 没有可比较的文本内容。")
    if len(chunking.chunks) == 1 and not chunking.truncated:
        return f"[完整原文]\n{chunking.chunks[0]}", chunking

    representations = []
    total_chunks = len(chunking.chunks)
    extraction_prompt = (
        "请把当前文本块整理为用于版本对比的结构化中间表示，简洁列出："
        "主题/章节、核心事实与条款、实体、数字、日期、金额、义务、例外和风险。"
        "不得推测，不得省略关键限定条件。"
    )
    for index, chunk in enumerate(chunking.chunks, start=1):
        if status_callback:
            status_callback(index, total_chunks)
        messages = [
            {"role": "system", "content": "你是严谨的文档结构化提取助手。"},
            {
                "role": "user",
                "content": (
                    f"{extraction_prompt}\n文件：{file_name}\n"
                    f"文本块：{index}/{total_chunks}\n\n{chunk}"
                ),
            },
        ]
        representation = complete_reply(messages, max_tokens=1200)
        representations.append(f"[文本块 {index}]\n{representation}")
    return "\n\n".join(representations), chunking


def to_xlsx_bytes(all_results, operation):
    import openpyxl
    from openpyxl.styles import Alignment, Font, PatternFill

    workbook = openpyxl.Workbook()
    worksheet = workbook.active
    worksheet.title = "AI处理汇总"
    for column, heading in enumerate(["文件名", "操作", "处理结果"], 1):
        cell = worksheet.cell(row=1, column=column, value=heading)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="4F6EF7")
        cell.alignment = Alignment(horizontal="center")
    for row, result in enumerate(all_results, 2):
        worksheet.cell(row=row, column=1, value=result["file"])
        worksheet.cell(row=row, column=2, value=operation)
        worksheet.cell(row=row, column=3, value=clean_md(result["result"]))
    worksheet.column_dimensions["A"].width = 30
    worksheet.column_dimensions["B"].width = 20
    worksheet.column_dimensions["C"].width = 80
    buffer = io.BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()

# ─────────────────────────────────────────────
#  侧边栏（与v1保持一致，加退出登录）
# ─────────────────────────────────────────────
def render_sidebar():
    with st.sidebar:
        st.markdown("### 🧠 DocMind AI")
        if st.session_state.username:
            safe_username = html.escape(st.session_state.username)
            st.markdown(f"<small style='color:rgba(255,255,255,0.5)'>👤 {safe_username}</small>",
                        unsafe_allow_html=True)
        else:
            st.markdown("<small style='color:rgba(255,255,255,0.4)'>智能文档处理平台</small>",
                        unsafe_allow_html=True)
        st.divider()

        st.markdown("**⚙️ API 配置**")
        if SERVER_API_KEY:
            st.info("🔒 当前使用服务器配置的 AI 服务")
            st.caption("服务器 API Key 与可信 Base URL 已绑定，前端不可查看或修改。")
            st.caption(f"🔒 Server Model: {SERVER_MODEL}")
        else:
            st.caption("BYOK 模式：密钥仅保留在当前 Streamlit Session，不会写入历史或磁盘。")
            provider_options = ["DeepSeek", "OpenAI", "自定义 OpenAI-Compatible API（高级）"]
            selected_provider = st.selectbox(
                "Provider",
                provider_options,
                key="byok_provider",
            )
            provider_defaults = {
                "DeepSeek": (DEFAULT_BASE_URL, "deepseek-chat"),
                "OpenAI": ("https://api.openai.com/v1", "gpt-4o-mini"),
            }
            previous_provider = st.session_state.get("byok_provider_applied")
            if selected_provider != previous_provider:
                if selected_provider in provider_defaults:
                    default_url, default_model = provider_defaults[selected_provider]
                    st.session_state.byok_base_url = default_url
                    st.session_state.model = default_model
                st.session_state.byok_provider_applied = selected_provider
            st.text_input(
                "API Key",
                type="password",
                placeholder="sk-...",
                key="byok_api_key",
            )
            st.text_input(
                "Base URL",
                placeholder=DEFAULT_BASE_URL,
                key="byok_base_url",
            )
            st.text_input(
                "Model",
                placeholder="deepseek-chat",
                key="model",
            )

        if st.button("🔌 测试连接", width="stretch"):
            try:
                get_active_ai_config()
            except AIConfigurationError as exc:
                st.error(str(exc))
            else:
                with st.spinner("测试中..."):
                    try:
                        complete_reply(
                            [{"role": "user", "content": "Reply with OK."}],
                            max_tokens=10,
                            track_usage=False,
                        )
                        st.success("✅ 连接成功")
                    except Exception as exc:
                        log_failure("AI connection test", exc)
                        st.error(safe_error_message("测试连接时", exc))

        st.divider()
        st.markdown("**🌐 联网搜索**")
        st.session_state.web_enabled = st.toggle("启用联网搜索",
                                                   value=st.session_state.web_enabled)
        if st.session_state.web_enabled:
            mode = st.selectbox("网络模式", ["auto","china","global"],
                                 index=["auto","china","global"].index(st.session_state.network_mode))
            st.session_state.network_mode = mode
            if SERVER_BOCHA_KEY:
                st.caption("🔒 博查使用服务器配置")
            else:
                st.text_input(
                    "博查 Key（国内优先）",
                    type="password",
                    placeholder="可留空，自动用百度",
                    key="byok_bocha_key",
                )
            if SERVER_TAVILY_KEY:
                st.caption("🔒 Tavily 使用服务器配置")
            else:
                st.text_input(
                    "Tavily Key（国外优先）",
                    type="password",
                    placeholder="可留空，自动用 DuckDuckGo",
                    key="byok_tavily_key",
                )

        st.divider()
        st.markdown("**📊 本次统计**")
        col1, col2 = st.columns(2)
        with col1:
            st.markdown(f"""<div class='stat-box'>
                <div class='stat-number'>{st.session_state.total_calls}</div>
                <div class='stat-label'>调用次数</div></div>""", unsafe_allow_html=True)
        with col2:
            tok = st.session_state.total_tokens
            tok_str = f"{tok//1000}k" if tok >= 1000 else str(tok)
            token_label = (
                "实际 Tokens"
                if st.session_state.actual_usage_calls and not st.session_state.estimated_calls
                else "估算 Tokens"
            )
            st.markdown(f"""<div class='stat-box'>
                <div class='stat-number'>{tok_str}</div>
                <div class='stat-label'>{token_label}</div></div>""", unsafe_allow_html=True)

        st.divider()
        # 手机端提示
        st.markdown("""
        <div class='mobile-key-hint' style='background:rgba(255,255,255,0.08);
             border-radius:8px;padding:10px 12px;font-size:0.78rem;
             color:rgba(255,255,255,0.7);line-height:1.5'>
          📱 手机用户：如当前为 BYOK 模式，请在此填入 API Key 后关闭侧边栏即可使用
        </div>""", unsafe_allow_html=True)

        if LOGIN_ON and st.session_state.logged_in:
            st.divider()
            if st.button("🚪 退出登录", width="stretch"):
                st.session_state.logged_in = False
                st.session_state.username  = ""
                st.rerun()

        st.markdown("<small style='color:rgba(255,255,255,0.3)'>DocMind AI v2.2 · Graceful Degradation</small>",
                    unsafe_allow_html=True)

# ─────────────────────────────────────────────
#  主页面品牌栏 + Tabs
# ─────────────────────────────────────────────
def render_main():
    username = st.session_state.username
    badge = f"👤 {html.escape(username)}" if username else "V2.2"
    st.markdown(f"""
    <div class='brand-bar'>
      <div class='brand-inner'>
        <div class='brand-logo'>🧠</div>
        <div>
          <p class='brand-title'>DocMind AI</p>
          <p class='brand-sub'>智能文档处理平台 · AI-Powered Document Intelligence</p>
        </div>
        <div class='brand-badge'>{badge}</div>
      </div>
    </div>
    """, unsafe_allow_html=True)

    # 未配置 AI 时保留所有基础能力，仅单独禁用 AI 增强。
    if not has_ai_config():
        st.markdown("""
        <div class='warn-card'>
          当前处于基础模式。未配置 AI 时仍可使用文件解析、文本提取、基础统计、文件对比和导出功能。
          配置 AI 后可解锁智能摘要、翻译、改写、风险分析和 AI 对话。
        </div>
        """, unsafe_allow_html=True)

    tab_doc, tab_cmp, tab_chat, tab_search, tab_history = st.tabs(
        ["📄 文档处理", "🔍 文件对比", "💬 AI 对话", "🌐 联网搜索", "📋 历史记录"])

    with tab_doc:     render_doc_tab()
    with tab_cmp:     render_compare_tab()
    with tab_chat:    render_chat_tab()
    with tab_search:  render_search_tab()
    with tab_history: render_history_tab()

# ─────────────────────────────────────────────
#  Tab 1：文档处理（与v1保持一致）
# ─────────────────────────────────────────────
OPERATIONS = {
    "📝 智能摘要":     {"desc":"提取核心内容",  "mode":"aggregate", "prompt":"请对以下文件内容进行智能摘要，提取核心观点、关键数据和重要结论，输出结构化的摘要报告："},
    "🌐 翻译成英文":   {"desc":"中译英",        "mode":"concatenate", "prompt":"请将以下内容完整、准确地翻译为英文，保持原文逻辑和格式，专业术语准确翻译："},
    "🇨🇳 翻译成中文":  {"desc":"英译中",        "mode":"concatenate", "prompt":"请将以下内容完整、准确地翻译为中文，保持原文逻辑和格式，语言自然流畅："},
    "⚠️ 风险分析":    {"desc":"合同/报告风险",  "mode":"aggregate", "prompt":"请仔细阅读以下内容，识别并分析所有潜在风险点、不合理条款和需要注意的事项，按风险等级（高/中/低）逐条列出："},
    "🔑 关键信息提取": {"desc":"结构化提取",    "mode":"aggregate", "prompt":"请从以下内容中提取所有关键信息，包括：时间、人名/机构、金额、核心条款、重要数据等，用结构化方式输出："},
    "✍️ 改写润色":    {"desc":"专业风格",       "mode":"concatenate", "prompt":"请将以下内容改写成正式、专业的风格，优化表达，消除语病，使其更适合商业/公文场景："},
    "📊 数据分析":     {"desc":"Excel/CSV专用", "mode":"aggregate", "prompt":"你是数据分析专家，请分析以下数据：统计规律、关键指标、异常值、趋势和业务洞察，给出具体建议："},
    "💡 自定义":       {"desc":"自由输入指令",  "mode":"aggregate", "prompt":""},
}


def render_basic_document(parsed, key_prefix: str):
    """展示完全不依赖 AI 的解析结果、统计与导出。"""
    with st.expander(f"📄 {parsed.file_name} · 基础解析", expanded=True):
        stats = parsed.stats
        metrics = [
            ("文件类型", parsed.file_type),
            ("字符数", f"{stats.get('characters', 0):,}"),
            ("段落数", f"{stats.get('paragraphs', 0):,}"),
        ]
        if "pages" in stats:
            metrics.append(("PDF 页数", f"{stats['pages']:,}"))
        elif "sheets" in stats:
            metrics.append(("Sheet 数", f"{stats['sheets']:,}"))
        elif "rows" in stats:
            metrics.append(("行数", f"{stats['rows']:,}"))
        else:
            metrics.append(("文件大小", f"{parsed.size_bytes / 1024:.1f} KB"))
        columns = st.columns(4)
        for column, (label, value) in zip(columns, metrics[:4]):
            column.metric(label, value)

        details = []
        if "rows" in stats:
            details.extend(
                [
                    f"行数 {stats['rows']:,}",
                    f"列数 {stats.get('columns', 0):,}",
                    f"非空单元格 {stats.get('non_empty_cells', 0):,}",
                    f"空值 {stats.get('empty_cells', 0):,}",
                ]
            )
        if stats.get("sheet_names"):
            details.append("Sheet：" + "、".join(map(str, stats["sheet_names"])))
        if details:
            st.caption(" · ".join(details))

        preview_limit = 12_000
        preview = parsed.text[:preview_limit]
        st.text_area(
            "提取文本预览",
            value=preview,
            height=220,
            disabled=True,
            key=f"preview_{key_prefix}",
        )
        if len(parsed.text) > preview_limit:
            st.caption(
                f"预览显示前 {preview_limit:,} 字符；下载文件包含全部 {len(parsed.text):,} 字符。"
            )

        download_txt, download_word = st.columns(2)
        with download_txt:
            st.download_button(
                "⬇️ 下载提取文本",
                data=parsed.text.encode("utf-8-sig"),
                file_name=f"{Path(parsed.file_name).stem}_提取文本.txt",
                mime="text/plain",
                key=f"extract_txt_{key_prefix}",
                width="stretch",
            )
        with download_word:
            if len(parsed.text) > MAX_BASIC_WORD_EXPORT_CHARS:
                st.caption("正文较长，为控制内存占用仅提供完整 TXT 导出。")
            else:
                try:
                    word_data = to_docx_bytes(parsed.text, f"提取文本：{parsed.file_name}")
                    st.download_button(
                        "⬇️ 导出 Word",
                        data=word_data,
                        file_name=f"{Path(parsed.file_name).stem}_提取文本.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=f"extract_word_{key_prefix}",
                        width="stretch",
                    )
                except Exception as exc:
                    log_failure("export extracted Word", exc)
                    st.caption("Word 导出暂不可用，仍可下载提取文本。")


def render_doc_tab():
    ai_available = has_ai_config()
    if not ai_available:
        st.info(
            "当前处于基础模式。未配置 AI 时仍可使用文件解析、文本提取、基础统计、"
            "文件对比和导出功能。配置 AI 后可解锁智能摘要、翻译、改写、风险分析和 AI 对话。"
        )

    col_left, col_right = st.columns([1, 1.2], gap="large")

    with col_left:
        st.markdown("#### 📁 上传文件")
        uploaded = st.file_uploader(
            "支持 PDF / DOCX / XLSX / CSV / TXT",
            type=["pdf","docx","xlsx","csv","txt"],
            accept_multiple_files=True, label_visibility="collapsed")

        if uploaded:
            for f in uploaded:
                size_kb = len(f.getvalue()) // 1024
                safe_file_name = html.escape(f.name)
                st.markdown(f"""
                <div class='file-chip'>📄 {safe_file_name}
                  <span style='color:#8892a4;font-size:0.75rem'>{size_kb} KB</span>
                </div>""", unsafe_allow_html=True)

        st.markdown("#### ⚡ 选择操作")
        op_keys     = list(OPERATIONS.keys())
        selected_op = st.session_state.get("selected_op", op_keys[0])
        cols        = st.columns(4)
        for i, op in enumerate(op_keys):
            with cols[i % 4]:
                if st.button(op, key=f"op_{i}", width="stretch",
                             type="primary" if op == selected_op else "secondary"):
                    st.session_state.selected_op = op
                    selected_op = op

        if selected_op == "💡 自定义":
            custom_prompt = st.text_area("自定义指令", height=80,
                                          placeholder="请描述你想让 AI 做什么...")
        else:
            custom_prompt = ""
            st.markdown(f"""<div class='info-card'>
              💡 {OPERATIONS[selected_op]['desc']} · {OPERATIONS[selected_op]['prompt'][:60]}...
            </div>""", unsafe_allow_html=True)

        extra_q     = st.text_input("➕ 追加要求（可选）", placeholder="例：重点关注第三方责任条款")
        run_clicked = st.button("▶  开始处理", type="primary",
                                 width="stretch",
                                 disabled=not uploaded or not ai_available)
        if not ai_available:
            st.caption("🔒 AI 操作当前不可用；基础解析、统计与导出不受影响。")

    with col_right:
        st.markdown("#### 📊 处理结果")
        parsed_documents = []
        if uploaded:
            for file_index, uploaded_file in enumerate(uploaded):
                try:
                    parsed = read_uploaded_document(uploaded_file)
                    parsed_documents.append((uploaded_file, parsed))
                    file_digest = hashlib.sha256(uploaded_file.getvalue()).hexdigest()[:10]
                    render_basic_document(parsed, f"doc_{file_index}_{file_digest}")
                except Exception as exc:
                    log_failure(f"parse uploaded document {uploaded_file.name}", exc)
                    st.error(safe_error_message(f"解析 {uploaded_file.name} 时", exc))
        else:
            st.caption("上传文件后将自动显示正文预览、基础统计与下载入口。")

        if run_clicked and parsed_documents:
            prompt_base = custom_prompt if selected_op == "💡 自定义" \
                         else OPERATIONS[selected_op]["prompt"]
            if not prompt_base.strip():
                st.warning("请输入自定义指令后再开始处理。")
                return
            if extra_q:
                prompt_base += f"\n\n附加要求：{extra_q}"

            all_results = []
            progress    = st.progress(0, text="准备中...")

            for idx, (f, parsed) in enumerate(parsed_documents):
                progress.progress(idx / len(parsed_documents),
                                   text=f"处理 {idx+1}/{len(uploaded)}：{f.name}")
                safe_file_name = html.escape(f.name)
                st.markdown(f"**📄 {safe_file_name}**")
                ph = st.empty()
                try:
                    content = parsed.text
                    search_ctx = ""
                    if st.session_state.web_enabled:
                        with st.spinner("🌐 联网搜索相关信息..."):
                            search_outcome = web_search(f.name + " " + prompt_base[:30])
                        render_search_feedback(search_outcome)
                        if search_outcome.results:
                            search_ctx = "\n\n参考网络信息：\n" + "\n".join(
                                f"- {result['title']}: {result['content'][:200]}"
                                for result in search_outcome.results[:3]
                            )

                    def update_chunk(chunk_index, chunk_total):
                        overall = (idx + (chunk_index - 1) / max(chunk_total, 1)) / len(parsed_documents)
                        progress.progress(
                            overall,
                            text=(
                                f"处理 {idx+1}/{len(uploaded)}：{f.name} · "
                                f"文本块 {chunk_index}/{chunk_total}"
                            ),
                        )

                    ai_outcome = run_optional_ai(
                        lambda: process_long_document(
                            content=content,
                            file_name=f.name,
                            instruction=prompt_base,
                            operation_mode=OPERATIONS[selected_op]["mode"],
                            placeholder=ph,
                            search_context=search_ctx,
                            status_callback=update_chunk,
                        ),
                        context=f"process document {f.name}",
                        action=f"处理 {f.name} 时",
                    )
                    if not ai_outcome.succeeded:
                        st.warning(
                            "AI 服务暂时不可用，但文档已经成功解析。你仍可查看正文、"
                            "基础统计、文件差异并下载解析结果。"
                        )
                        st.caption(ai_outcome.error_message)
                        progress.progress(
                            (idx + 1) / len(parsed_documents),
                            text=f"已完成 {idx+1}/{len(parsed_documents)} 个文件",
                        )
                        continue
                    full, chunking, incomplete_output = ai_outcome.result
                    scope_notice = chunk_scope_notice(chunking)
                    if scope_notice:
                        st.warning(scope_notice)
                    elif len(chunking.chunks) > 1:
                        st.caption(f"✅ 已完成 {len(chunking.chunks)} 个文本块的综合处理。")
                    if incomplete_output:
                        st.warning(
                            "AI 已自动尝试续写，但输出仍达到长度上限；结果可能不完整，请分段重试。"
                        )
                    all_results.append({"file": f.name, "result": full})
                    save_history(selected_op, f.name, full)
                except Exception as exc:
                    log_failure(f"process document ({type(exc).__name__})", exc)
                    st.warning(
                        "AI 服务暂时不可用，但文档已经成功解析。你仍可查看正文、"
                        "基础统计、文件差异并下载解析结果。"
                    )
                    st.caption(safe_error_message(f"处理 {f.name} 时", exc))
                progress.progress(
                    (idx + 1) / len(parsed_documents),
                    text=f"已完成 {idx+1}/{len(parsed_documents)} 个文件",
                )

            progress.progress(1.0, text="✅ 全部完成！")

            if all_results:
                st.divider()
                combined = "\n\n".join(f"=== {r['file']} ===\n{clean_md(r['result'])}"
                                       for r in all_results)
                dc1, dc2, dc3 = st.columns(3)
                with dc1:
                    st.download_button("⬇️ 下载 TXT", data=combined.encode("utf-8"),
                        file_name=f"DocMind_{datetime.now().strftime('%m%d_%H%M')}.txt",
                        mime="text/plain", width="stretch")
                with dc2:
                    try:
                        st.download_button("⬇️ 下载 Word", data=to_docx_bytes(combined, selected_op),
                            file_name=f"DocMind_{datetime.now().strftime('%m%d_%H%M')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            width="stretch")
                    except Exception as exc:
                        log_failure("export Word", exc)
                        st.warning("Word 导出暂不可用，仍可下载 TXT。")
                with dc3:
                    if len(all_results) > 1:
                        try:
                            st.download_button("⬇️ 下载 Excel 汇总", data=to_xlsx_bytes(all_results, selected_op),
                                file_name=f"DocMind_汇总_{datetime.now().strftime('%m%d_%H%M')}.xlsx",
                                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                width="stretch")
                        except Exception as exc:
                            log_failure("export Excel", exc)
                            st.warning("Excel 汇总导出暂不可用，仍可下载 TXT。")

# ─────────────────────────────────────────────
#  Tab 2：文件对比分析（新增）
# ─────────────────────────────────────────────
def render_compare_tab():
    st.markdown("""<div class='info-card'>
      📌 基础精确 Diff 无需 AI；配置 AI 后可进一步解释关键变化、影响和风险。
    </div>""", unsafe_allow_html=True)

    col1, col2 = st.columns(2, gap="medium")
    with col1:
        st.markdown("**📄 文件 A（原版）**")
        file_a = st.file_uploader("文件A", type=["pdf","docx","xlsx","csv","txt"],
                                   label_visibility="collapsed", key="cmp_a")
        if file_a:
            st.markdown(f"<div class='file-chip'>✅ {html.escape(file_a.name)}</div>", unsafe_allow_html=True)
    with col2:
        st.markdown("**📄 文件 B（新版）**")
        file_b = st.file_uploader("文件B", type=["pdf","docx","xlsx","csv","txt"],
                                   label_visibility="collapsed", key="cmp_b")
        if file_b:
            st.markdown(f"<div class='file-chip'>✅ {html.escape(file_b.name)}</div>", unsafe_allow_html=True)

    cmp_mode = st.selectbox("对比维度", [
        "全面对比（找出所有差异）",
        "条款变更（合同专用）",
        "数据变化（Excel/报表专用）",
        "风险变化（新增/消除的风险点）",
    ])
    focus = st.text_input("重点关注（可选）", placeholder="例：付款条款、违约责任、截止日期")

    current_signature = None
    if file_a and file_b:
        current_signature = hashlib.sha256(
            file_a.getvalue() + b"\0" + file_b.getvalue()
        ).hexdigest()
    compare_state = st.session_state.compare_state
    if compare_state and compare_state.get("signature") != current_signature:
        st.session_state.compare_state = None
        compare_state = None

    if st.button("🔍 开始基础精确对比", type="primary", disabled=not (file_a and file_b)):
        try:
            parsed_a = read_uploaded_document(file_a)
            parsed_b = read_uploaded_document(file_b)
            diff = basic_text_diff(
                parsed_a.text,
                parsed_b.text,
                file_a.name,
                file_b.name,
            )
            compare_state = {
                "signature": current_signature,
                "parsed_a": parsed_a,
                "parsed_b": parsed_b,
                "diff": diff,
                "ai_result": "",
            }
            st.session_state.compare_state = compare_state
        except Exception as exc:
            log_failure("basic document diff", exc)
            st.error(safe_error_message("执行基础对比时", exc))

    compare_state = st.session_state.compare_state
    if not compare_state:
        if not has_ai_config():
            st.caption("当前为基础模式：上传两份文件即可使用精确 Diff；AI 解释按钮将保持禁用。")
        return

    parsed_a = compare_state["parsed_a"]
    parsed_b = compare_state["parsed_b"]
    diff = compare_state["diff"]
    st.divider()
    st.markdown(
        f"**🔍 基础对比：`{html.escape(parsed_a.file_name)}` vs `{html.escape(parsed_b.file_name)}`**"
    )
    metric_columns = st.columns(5)
    metric_values = [
        ("新增行", len(diff.added)),
        ("删除行", len(diff.removed)),
        ("修改行", len(diff.modified)),
        ("未变行", diff.unchanged_lines),
        ("相似度", f"{diff.similarity:.1%}"),
    ]
    for column, (label, value) in zip(metric_columns, metric_values):
        column.metric(label, value)

    if not diff.change_count:
        st.success("两份文件的提取文本完全一致。")
    else:
        preview_lines = 300
        with st.expander(f"➕ 新增内容（{len(diff.added)} 行）"):
            st.code("\n".join(diff.added[:preview_lines]) or "无", language=None)
            if len(diff.added) > preview_lines:
                st.caption("页面仅显示前 300 行；下载报告包含完整 Diff。")
        with st.expander(f"➖ 删除内容（{len(diff.removed)} 行）"):
            st.code("\n".join(diff.removed[:preview_lines]) or "无", language=None)
            if len(diff.removed) > preview_lines:
                st.caption("页面仅显示前 300 行；下载报告包含完整 Diff。")
        with st.expander(f"✏️ 修改内容（{len(diff.modified)} 处）"):
            modified_preview = "\n\n".join(
                f"- {item['before']}\n+ {item['after']}"
                for item in diff.modified[:preview_lines]
            )
            st.code(modified_preview or "无", language="diff")
            if len(diff.modified) > preview_lines:
                st.caption("页面仅显示前 300 处；下载报告包含完整 Diff。")
        with st.expander("完整 Unified Diff"):
            st.code(diff.unified_diff[:30_000] or "无差异", language="diff")
            if len(diff.unified_diff) > 30_000:
                st.caption("页面预览已限制为 30,000 字符；下载报告包含完整 Diff。")

    basic_report = (
        f"DocMind 基础精确 Diff\n"
        f"文件 A：{parsed_a.file_name}\n文件 B：{parsed_b.file_name}\n"
        f"新增行：{len(diff.added)}\n删除行：{len(diff.removed)}\n"
        f"修改行：{len(diff.modified)}\n相似度：{diff.similarity:.2%}\n\n"
        f"{diff.unified_diff or '无差异'}"
    )
    st.download_button(
        "⬇️ 下载基础 Diff 报告",
        data=basic_report.encode("utf-8-sig"),
        file_name=f"基础Diff_{datetime.now().strftime('%m%d_%H%M')}.txt",
        mime="text/plain",
        width="stretch",
    )

    ai_available = has_ai_config()
    ai_clicked = st.button(
        "🧠 使用 AI 解释主要变化",
        type="primary",
        disabled=not ai_available,
        width="stretch",
    )
    if not ai_available:
        st.info("AI 解释当前不可用；上方基础 Diff、差异统计与下载功能仍可正常使用。")

    if ai_clicked:
        progress = st.progress(0, text="准备 AI 解释...")
        try:
            content_a = parsed_a.text
            content_b = parsed_b.text

            def update_a(chunk_index, chunk_total):
                progress.progress(
                    0.05 + 0.4 * chunk_index / chunk_total,
                    text=f"结构化文件 A：文本块 {chunk_index}/{chunk_total}",
                )

            def update_b(chunk_index, chunk_total):
                progress.progress(
                    0.5 + 0.35 * chunk_index / chunk_total,
                    text=f"结构化文件 B：文本块 {chunk_index}/{chunk_total}",
                )

            comparison_a, chunking_a = prepare_comparison_content(
                parsed_a.file_name, content_a, update_a
            )
            comparison_b, chunking_b = prepare_comparison_content(
                parsed_b.file_name, content_b, update_b
            )
            for label, chunking in (("文件 A", chunking_a), ("文件 B", chunking_b)):
                notice = chunk_scope_notice(chunking)
                if notice:
                    st.warning(f"{label}：{notice}")
                elif len(chunking.chunks) > 1:
                    st.caption(f"{label}：已覆盖 {len(chunking.chunks)} 个文本块。")

            prompt = f"""你是专业文档对比分析师。请对以下两份文件进行{cmp_mode}。
{"重点关注：" + focus if focus else ""}

请严格按以下结构输出，并以原文或中间表示为依据：
1. 综合结论与总体变化摘要
2. 核心相同点
3. 新增内容（B 有、A 没有）
4. 删除内容（A 有、B 删除）
5. 修改内容（含表述、义务、范围和例外条件变化）
6. 数字、日期、金额等关键变化
7. 风险变化（新增、加重、降低或消除）

---文件 A：{parsed_a.file_name}---
{comparison_a}

---文件 B：{parsed_b.file_name}---
{comparison_b}"""

            messages = [
                {"role":"system","content":"你是严谨的文档对比分析师。区分相同、新增、删除和修改，不确定时明确说明，不得编造。"},
                {"role":"user","content":prompt}
            ]

            progress.progress(0.9, text="综合两份文档的差异...")
            full = complete_reply(messages)
            compare_state["ai_result"] = full
            st.session_state.compare_state = compare_state
            save_history(
                "AI 对比解释",
                f"{parsed_a.file_name} vs {parsed_b.file_name}",
                full,
            )
            progress.progress(1.0, text="✅ 对比完成")
        except Exception as exc:
            log_failure("AI comparison explanation", exc)
            st.warning(
                "AI 服务暂时不可用，但基础 Diff 已完成。你仍可查看差异统计并下载基础报告。"
            )
            st.caption(safe_error_message("解释文件差异时", exc))

    ai_result = compare_state.get("ai_result", "")
    if ai_result:
        st.divider()
        st.markdown("**🧠 AI 变化解释**")
        st.markdown(clean_md(ai_result))
        dl1, dl2 = st.columns(2)
        with dl1:
            st.download_button(
                "⬇️ 下载 AI 解释 TXT",
                data=clean_md(ai_result).encode("utf-8-sig"),
                file_name=f"AI对比解释_{datetime.now().strftime('%m%d_%H%M')}.txt",
                mime="text/plain",
                width="stretch",
            )
        with dl2:
            try:
                st.download_button(
                    "⬇️ 下载 AI 解释 Word",
                    data=to_docx_bytes(
                        ai_result,
                        f"文件对比：{parsed_a.file_name} vs {parsed_b.file_name}",
                    ),
                    file_name=f"AI对比解释_{datetime.now().strftime('%m%d_%H%M')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    width="stretch",
                )
            except Exception as exc:
                log_failure("export comparison Word", exc)
                st.warning("Word 导出暂不可用，仍可下载 TXT。")

# ─────────────────────────────────────────────
#  Tab 3：AI 对话（与v1保持一致）
# ─────────────────────────────────────────────
def render_chat_tab():
    if not has_ai_config():
        st.info("👈 请先在左侧侧边栏填入 API Key")
        return

    for msg in st.session_state.chat_messages:
        with st.chat_message(msg["role"], avatar="🧑" if msg["role"]=="user" else "🧠"):
            st.markdown(clean_md(msg["content"]))

    if prompt := st.chat_input("输入你的问题，或者粘贴文字让 AI 分析..."):
        st.session_state.chat_messages.append({"role":"user","content":prompt})
        with st.chat_message("user", avatar="🧑"):
            st.markdown(prompt)
        with st.chat_message("assistant", avatar="🧠"):
            ph = st.empty()
            search_ctx = ""
            if st.session_state.web_enabled:
                with st.spinner("🌐 搜索中..."):
                    search_outcome = web_search(prompt)
                    render_search_feedback(search_outcome)
                    if search_outcome.results:
                        search_ctx = "\n\n参考网络最新信息：\n" + \
                            "\n".join(
                                f"[{result['engine']}] {result['title']}: {result['content'][:200]}"
                                for result in search_outcome.results[:4]
                            )
            messages = [
                {"role":"system","content":"你是专业智能助手，回答清晰、简洁、有条理。不要使用过多Markdown符号。"},
            ] + st.session_state.chat_messages[-19:-1] + [
                {"role":"user","content": prompt + search_ctx}
            ]
            try:
                full = stream_reply(messages, ph)
                st.session_state.chat_messages.append({"role":"assistant","content":full})
            except Exception as exc:
                log_failure("AI chat", exc)
                st.error(safe_error_message("生成回复时", exc))

    if st.session_state.chat_messages:
        if st.button("🗑️ 清空对话", key="clear_chat"):
            st.session_state.chat_messages = []
            st.rerun()

# ─────────────────────────────────────────────
#  Tab 4：联网搜索（与v1保持一致）
# ─────────────────────────────────────────────
def render_search_tab():
    st.markdown("#### 🌐 联网搜索")
    st.markdown("<small style='color:#8892a4'>自动检测网络环境：国内用博查/百度，国外用Tavily/DuckDuckGo</small>",
                unsafe_allow_html=True)
    col1, col2 = st.columns([3, 1])
    with col1:
        query = st.text_input("搜索关键词", placeholder="输入要搜索的内容...", label_visibility="collapsed")
    with col2:
        search_btn = st.button("🔍 搜索", width="stretch", type="primary")

    if search_btn and query:
        with st.spinner("搜索中..."):
            search_outcome = web_search(query)
        render_search_feedback(search_outcome)
        results = search_outcome.results
        if results:
            engine = results[0].get("engine","")
            st.markdown(f"<span class='tag'>引擎：{engine}</span>"
                        f"<span class='tag'>找到 {len(results)} 条结果</span>",
                        unsafe_allow_html=True)
            for r in results:
                safe_title = html.escape(r["title"])
                safe_url = html.escape(r["url"][:80])
                safe_content = html.escape(r["content"] or "暂无摘要")
                st.markdown(f"""<div class='search-item'>
                  <div class='search-title'>{safe_title}</div>
                  <div class='search-url'>{safe_url}</div>
                  <div class='search-snip'>{safe_content}</div>
                </div>""", unsafe_allow_html=True)
            if has_ai_config():
                if st.button("🧠 让 AI 综合分析这些结果"):
                    ctx = "\n".join(f"[{r['title']}] {r['content']}" for r in results)
                    with st.spinner("AI 分析中..."):
                        try:
                            messages = [
                                {"role":"system","content":"你是信息综合分析专家"},
                                {"role":"user","content":f"请综合以下搜索结果，回答问题「{query}」，给出清晰的分析结论：\n\n{ctx}"}
                            ]
                            full = complete_reply(messages)
                            st.markdown("**AI 综合分析：**")
                            st.markdown(clean_md(full))
                        except Exception as exc:
                            log_failure("AI search synthesis", exc)
                            st.error(safe_error_message("综合搜索结果时", exc))
        elif not search_outcome.error:
            st.warning("未找到结果，请检查网络或更换关键词")

# ─────────────────────────────────────────────
#  Tab 5：历史记录（与v1保持一致）
# ─────────────────────────────────────────────
def render_history_tab():
    st.markdown("#### 📋 处理历史")
    if not st.session_state.history:
        st.info("暂无处理记录，处理文件后会自动记录在这里")
        return
    for i, h in enumerate(st.session_state.history[:20]):
        with st.expander(f"[{h['time']}] {h['op']} · {h['file']}", expanded=(i==0)):
            st.markdown(clean_md(h["result"]))
            st.download_button("⬇️ 下载", data=clean_md(h["result"]).encode("utf-8"),
                                file_name=f"{h['file']}_AI结果.txt",
                                key=f"dl_h_{i}", mime="text/plain")
    if st.button("🗑️ 清空历史"):
        st.session_state.history = []
        st.rerun()

# ─────────────────────────────────────────────
#  主入口
# ─────────────────────────────────────────────
if LOGIN_CONFIG_ERROR:
    st.error("登录配置无效，请管理员检查 DOCMIND_USERS。")
    st.stop()

if LOGIN_ON and not st.session_state.logged_in:
    render_login()
    st.stop()

render_sidebar()
render_main()
