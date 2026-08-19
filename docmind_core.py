"""DocMind 的可测试核心逻辑：文件解析、分块、配置隔离与联网搜索。"""

from __future__ import annotations

import csv
import difflib
import html
import io
import ipaddress
import logging
import re
import socket
import urllib.parse
import urllib.request
import zipfile
from dataclasses import dataclass, field
from functools import lru_cache
from pathlib import Path
from typing import Any, Callable, Dict, List, Mapping, Optional, Sequence
from urllib.parse import urlparse


logger = logging.getLogger("docmind")

DEFAULT_BASE_URL = "https://api.deepseek.com"
DEFAULT_CHUNK_SIZE = 6500
DEFAULT_CHUNK_OVERLAP = 300
TRANSFORM_CHUNK_SIZE = 3500
MAX_CHUNKS = 10
MAX_TRANSFORM_CONTINUATIONS = 2
MAX_FILE_SIZE_BYTES = 40 * 1024 * 1024
MAX_PDF_PAGES = 300
MAX_XLSX_SHEETS = 50
MAX_XLSX_EFFECTIVE_CELLS = 500_000
MAX_XLSX_SCANNED_CELLS = 2_000_000
MAX_XLSX_UNCOMPRESSED_BYTES = 100 * 1024 * 1024
MAX_DOCX_UNCOMPRESSED_BYTES = 50 * 1024 * 1024
MAX_DOCX_TEXT_CHARS = 2_000_000
MAX_CSV_ROWS = 200_000
MAX_DIFF_LINES = 50_000
SUPPORTED_EXTENSIONS = (".pdf", ".docx", ".xlsx", ".csv", ".txt")


class DocMindError(Exception):
    """可安全展示给最终用户的 DocMind 业务异常。"""


class DocumentReadError(DocMindError):
    """文件无法解析或没有可处理内容。"""


class UnsupportedFileTypeError(DocumentReadError):
    """上传了未支持的文件类型。"""


class AIConfigurationError(DocMindError):
    """AI 服务配置缺失或无效。"""


class AIResponseError(DocMindError):
    """AI 服务未返回可用内容。"""


@dataclass(frozen=True)
class ChunkingResult:
    chunks: List[str]
    truncated: bool
    total_chars: int
    processed_chars: int


@dataclass(frozen=True)
class ParsedDocument:
    file_name: str
    file_type: str
    size_bytes: int
    text: str
    stats: Mapping[str, Any]


@dataclass(frozen=True)
class BasicDiffResult:
    added: List[str]
    removed: List[str]
    modified: List[Dict[str, str]]
    unified_diff: str
    similarity: float
    unchanged_lines: int

    @property
    def change_count(self) -> int:
        return len(self.added) + len(self.removed) + len(self.modified)


@dataclass(frozen=True)
class CompletionResult:
    text: str
    finish_reason: str

    @property
    def truncated(self) -> bool:
        return self.finish_reason == "length"


@dataclass(frozen=True)
class ContinuationResult:
    text: str
    incomplete: bool
    continuation_count: int


@dataclass(frozen=True)
class AIEnhancementOutcome:
    result: Any = None
    error_message: Optional[str] = None

    @property
    def succeeded(self) -> bool:
        return self.error_message is None


@dataclass(frozen=True, repr=False)
class AIConfig:
    api_key: str = field(repr=False)
    base_url: str
    model: str
    source: str

    def __repr__(self) -> str:
        return (
            "AIConfig(api_key=<redacted>, "
            f"base_url={self.base_url!r}, model={self.model!r}, source={self.source!r})"
        )


@dataclass(frozen=True)
class TokenUsage:
    prompt_tokens: int
    completion_tokens: int
    total_tokens: int
    estimated: bool


@dataclass(frozen=True)
class SearchOutcome:
    results: List[Dict[str, str]]
    notices: List[str] = field(default_factory=list)
    error: Optional[str] = None


def redact_sensitive(value: str) -> str:
    """移除日志文本中常见的密钥、Bearer Token 与 URL 凭证。"""
    text = str(value)
    patterns = (
        (r"(?i)(api[_-]?key\s*[=:]\s*)[^\s,;&]+", r"\1<redacted>"),
        (r"(?i)(authorization\s*[=:]\s*bearer\s+)[^\s,;&]+", r"\1<redacted>"),
        (r"(?i)(bearer\s+)[A-Za-z0-9._~+\-/=]+", r"\1<redacted>"),
        (r"\bsk-[A-Za-z0-9_-]{8,}\b", "<redacted-key>"),
        (r"(https?://)[^/@\s:]+:[^/@\s]+@", r"\1<redacted>@"),
    )
    for pattern, replacement in patterns:
        text = re.sub(pattern, replacement, text)
    return text[:800]


def log_failure(context: str, exc: BaseException) -> None:
    """记录可诊断但不含密钥的错误。"""
    logger.error(
        "%s failed [%s]: %s",
        redact_sensitive(context),
        type(exc).__name__,
        redact_sensitive(str(exc)),
    )


def safe_error_message(action: str, exc: BaseException) -> str:
    """生成不暴露服务器路径、环境变量或第三方响应细节的提示。"""
    if isinstance(exc, DocMindError):
        return f"操作失败：{exc}"

    name = type(exc).__name__.lower()
    status_code = getattr(exc, "status_code", None)
    response = getattr(exc, "response", None)
    if status_code is None and response is not None:
        status_code = getattr(response, "status_code", None)
    if status_code in (401, 403) or "authentication" in name or "permission" in name:
        detail = "AI 服务鉴权失败，请检查 API Key。"
    elif status_code == 402:
        detail = "AI 服务额度不足或账户欠费，请检查服务商账户。"
    elif status_code == 429 or "ratelimit" in name:
        detail = "AI 服务请求过于频繁或额度不足，请稍后重试。"
    elif isinstance(status_code, int) and status_code >= 500:
        detail = "AI 服务暂时异常，请稍后重试。"
    elif "timeout" in name:
        detail = "请求超时，请稍后重试或缩小处理范围。"
    elif "connection" in name or "network" in name:
        detail = "网络连接失败，请稍后重试。"
    elif "notfound" in name or status_code == 404:
        detail = "AI 模型或接口不存在，请检查 Base URL 与 Model。"
    else:
        detail = "暂时无法完成，请稍后重试。"
    return f"操作失败：{action}{detail}"


def run_optional_ai(
    operation: Callable[[], Any],
    *,
    context: str,
    action: str,
) -> AIEnhancementOutcome:
    """执行可选 AI 增强；失败时返回安全提示而不影响已解析的基础结果。"""
    try:
        return AIEnhancementOutcome(result=operation())
    except Exception as exc:
        log_failure(context, exc)
        return AIEnhancementOutcome(error_message=safe_error_message(action, exc))


def make_history_entry(
    operation: str,
    file_name: str,
    result: str,
    time_label: str,
) -> Dict[str, str]:
    """构建只含处理结果的会话历史项，不接受或存储 AI 配置。"""
    return {
        "time": str(time_label),
        "op": str(operation),
        "file": str(file_name),
        "result": str(result),
    }


def _decode_text(data: bytes, file_label: str) -> str:
    if not data:
        raise DocumentReadError(f"{file_label} 是空文件。")

    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise DocumentReadError(
        f"{file_label} 的文本编码无法识别，请另存为 UTF-8 或 GB18030 后重试。"
    )


def _text_stats(text: str) -> Dict[str, int]:
    lines = text.splitlines()
    paragraphs = [part for part in re.split(r"\n\s*\n", text) if part.strip()]
    return {
        "characters": len(text),
        "non_whitespace_characters": len(re.sub(r"\s", "", text)),
        "lines": len(lines),
        "paragraphs": len(paragraphs),
    }


def _inspect_ooxml_archive(
    data: bytes,
    file_name: str,
    max_uncompressed_bytes: int,
) -> None:
    """在交给 OOXML 解析器前限制 ZIP 解压规模。"""
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as archive:
            total_uncompressed = sum(info.file_size for info in archive.infolist())
            if total_uncompressed > max_uncompressed_bytes:
                raise DocumentReadError(
                    f"{file_name} 解压后的内容超过 {max_uncompressed_bytes // 1024 // 1024} MB 限制。"
                )
    except DocumentReadError:
        raise
    except (zipfile.BadZipFile, OSError) as exc:
        raise DocumentReadError(f"无法读取 {file_name}，OOXML 文件结构无效。") from exc


def parse_document(file_name: str, data: bytes) -> ParsedDocument:
    """安全解析上传文件，并返回文本和无需 AI 的基础统计。"""
    suffix = Path(file_name).suffix.lower()
    if suffix == ".xls":
        raise UnsupportedFileTypeError("不支持旧版 .xls，请先转换为 .xlsx。")
    if suffix not in SUPPORTED_EXTENSIONS:
        allowed = "、".join(ext.lstrip(".").upper() for ext in SUPPORTED_EXTENSIONS)
        raise UnsupportedFileTypeError(f"不支持 {suffix or '无扩展名'} 文件；支持格式：{allowed}。")
    if not data:
        raise DocumentReadError(f"{file_name} 是空文件。")
    if len(data) > MAX_FILE_SIZE_BYTES:
        raise DocumentReadError(
            f"{file_name} 超过应用层单文件 {MAX_FILE_SIZE_BYTES // 1024 // 1024} MB 限制。"
        )

    try:
        stats: Dict[str, Any] = {}
        if suffix == ".txt":
            text = _decode_text(data, file_name)
        elif suffix == ".csv":
            text = _decode_text(data, file_name)
            row_count = 0
            max_columns = 0
            non_empty_cells = 0
            empty_cells = 0
            for row_count, row in enumerate(csv.reader(io.StringIO(text)), start=1):
                if row_count > MAX_CSV_ROWS:
                    raise DocumentReadError(
                        f"{file_name} 超过 CSV 最大 {MAX_CSV_ROWS:,} 行限制。"
                    )
                max_columns = max(max_columns, len(row))
                non_empty_cells += sum(bool(value.strip()) for value in row)
                empty_cells += sum(not value.strip() for value in row)
            stats.update(
                {
                    "rows": row_count,
                    "columns": max_columns,
                    "non_empty_cells": non_empty_cells,
                    "empty_cells": empty_cells,
                }
            )
        elif suffix == ".docx":
            import docx

            _inspect_ooxml_archive(data, file_name, MAX_DOCX_UNCOMPRESSED_BYTES)
            document = docx.Document(io.BytesIO(data))
            lines = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
            content_paragraphs = len(lines)
            table_rows = 0
            for table in document.tables:
                for row in table.rows:
                    table_rows += 1
                    lines.append("\t".join(cell.text for cell in row.cells))
            text = "\n".join(lines)
            if len(text) > MAX_DOCX_TEXT_CHARS:
                raise DocumentReadError(
                    f"{file_name} 提取文本超过 {MAX_DOCX_TEXT_CHARS:,} 字符限制。"
                )
            stats.update(
                {
                    "document_paragraphs": len(document.paragraphs),
                    "content_paragraphs": content_paragraphs,
                    "tables": len(document.tables),
                    "table_rows": table_rows,
                }
            )
        elif suffix == ".xlsx":
            import openpyxl

            _inspect_ooxml_archive(data, file_name, MAX_XLSX_UNCOMPRESSED_BYTES)
            workbook = openpyxl.load_workbook(
                io.BytesIO(data), read_only=True, data_only=True
            )
            if len(workbook.worksheets) > MAX_XLSX_SHEETS:
                workbook.close()
                raise DocumentReadError(
                    f"{file_name} 超过 XLSX 最大 {MAX_XLSX_SHEETS} 个 Sheet 限制。"
                )
            lines = []
            effective_cells = 0
            scanned_cells = 0
            total_rows = 0
            max_columns = 0
            sheet_names = []
            for worksheet in workbook.worksheets:
                sheet_names.append(worksheet.title)
                lines.append(f"[Sheet: {worksheet.title}]")
                declared_cells = (worksheet.max_row or 0) * (worksheet.max_column or 0)
                if declared_cells > MAX_XLSX_SCANNED_CELLS:
                    workbook.close()
                    raise DocumentReadError(
                        f"{file_name} 工作表扫描范围过大，请删除多余格式或空白区域后重试。"
                    )
                for row in worksheet.iter_rows(values_only=True):
                    scanned_cells += len(row)
                    if scanned_cells > MAX_XLSX_SCANNED_CELLS:
                        workbook.close()
                        raise DocumentReadError(
                            f"{file_name} 工作表扫描范围过大，请删除多余格式或空白区域后重试。"
                        )
                    values = [str(cell) if cell is not None else "" for cell in row]
                    if any(values):
                        total_rows += 1
                        last_value = max(index for index, value in enumerate(values, 1) if value)
                        max_columns = max(max_columns, last_value)
                        effective_cells += sum(bool(value) for value in values)
                        if effective_cells > MAX_XLSX_EFFECTIVE_CELLS:
                            workbook.close()
                            raise DocumentReadError(
                                f"{file_name} 超过 XLSX 最大 {MAX_XLSX_EFFECTIVE_CELLS:,} 个有效单元格限制。"
                            )
                        lines.append("\t".join(values))
            workbook.close()
            if not effective_cells:
                raise DocumentReadError(f"{file_name} 没有可处理的单元格内容。")
            text = "\n".join(lines)
            stats.update(
                {
                    "sheets": len(sheet_names),
                    "sheet_names": sheet_names,
                    "rows": total_rows,
                    "columns": max_columns,
                    "non_empty_cells": effective_cells,
                    "empty_cells": max(total_rows * max_columns - effective_cells, 0),
                }
            )
        else:
            import pdfplumber

            with pdfplumber.open(io.BytesIO(data)) as pdf:
                if len(pdf.pages) > MAX_PDF_PAGES:
                    raise DocumentReadError(
                        f"{file_name} 超过 PDF 最大 {MAX_PDF_PAGES} 页限制。"
                    )
                pages = [page.extract_text() or "" for page in pdf.pages]
            text = "\n\n".join(pages)
            stats.update({"pages": len(pages), "pages_with_text": sum(bool(page.strip()) for page in pages)})
    except DocumentReadError:
        raise
    except Exception as exc:
        log_failure(f"parse {suffix or 'unknown'} document", exc)
        raise DocumentReadError(
            f"无法读取 {file_name}，文件可能已损坏或格式不正确。"
        ) from exc

    text = text.strip()
    if not text:
        if suffix == ".pdf":
            raise DocumentReadError(
                f"{file_name} 没有可提取的文字；扫描型 PDF 需要先进行 OCR。"
            )
        raise DocumentReadError(f"{file_name} 没有可处理的文本内容。")
    stats.update(_text_stats(text))
    if suffix == ".docx":
        stats["paragraphs"] = stats.get("content_paragraphs", stats["paragraphs"])
    return ParsedDocument(
        file_name=file_name,
        file_type=suffix.lstrip(".").upper(),
        size_bytes=len(data),
        text=text,
        stats=stats,
    )


def parse_document_bytes(file_name: str, data: bytes) -> str:
    """兼容旧调用：只返回安全解析后的文本。"""
    return parse_document(file_name, data).text


def split_text_into_chunks(
    text: str,
    chunk_size: int = DEFAULT_CHUNK_SIZE,
    overlap: int = DEFAULT_CHUNK_OVERLAP,
    max_chunks: int = MAX_CHUNKS,
) -> ChunkingResult:
    """优先在段落或句子边界切分文本，并限制最多处理块数。"""
    if chunk_size < 500:
        raise ValueError("chunk_size 必须至少为 500。")
    if overlap < 0 or overlap >= chunk_size:
        raise ValueError("overlap 必须大于等于 0 且小于 chunk_size。")
    if max_chunks < 1:
        raise ValueError("max_chunks 必须至少为 1。")

    normalized = str(text or "").replace("\r\n", "\n").replace("\r", "\n").strip()
    if not normalized:
        return ChunkingResult([], False, 0, 0)

    chunks: List[str] = []
    start = 0
    last_unique_end = 0
    text_length = len(normalized)

    while start < text_length and len(chunks) < max_chunks:
        hard_end = min(start + chunk_size, text_length)
        end = hard_end

        if hard_end < text_length:
            window = normalized[start:hard_end]
            minimum = int(chunk_size * 0.55)
            candidates = [window.rfind("\n\n"), window.rfind("\n")]
            sentence_matches = list(re.finditer(r"[。！？；.!?]\s*", window))
            if sentence_matches:
                candidates.append(sentence_matches[-1].end())
            viable = [position for position in candidates if position >= minimum]
            if viable:
                end = start + max(viable)

        chunk = normalized[start:end].strip()
        if chunk:
            chunks.append(chunk)
        last_unique_end = max(last_unique_end, end)
        if end >= text_length:
            break

        next_start = max(end - overlap, start + 1)
        while next_start < end and normalized[next_start].isspace():
            next_start += 1
        start = next_start

    truncated = last_unique_end < text_length
    return ChunkingResult(chunks, truncated, text_length, last_unique_end)


def chunk_scope_notice(result: ChunkingResult) -> Optional[str]:
    if not result.truncated:
        return None
    return (
        f"文档较长，本次最多处理前 {len(result.chunks)} 个文本块"
        f"（约 {result.processed_chars:,}/{result.total_chars:,} 字符）。"
    )


def basic_text_diff(
    original: str,
    revised: str,
    original_name: str = "文件 A",
    revised_name: str = "文件 B",
) -> BasicDiffResult:
    """生成无需 AI 的逐行精确 Diff，并把 replace 与增删分开统计。"""
    original_lines = str(original or "").splitlines()
    revised_lines = str(revised or "").splitlines()
    if len(original_lines) + len(revised_lines) > MAX_DIFF_LINES:
        raise DocumentReadError(
            f"两份文件合计超过基础 Diff 最大 {MAX_DIFF_LINES:,} 行限制。"
        )

    matcher = difflib.SequenceMatcher(None, original_lines, revised_lines, autojunk=False)
    added: List[str] = []
    removed: List[str] = []
    modified: List[Dict[str, str]] = []
    unchanged_lines = 0

    for tag, a_start, a_end, b_start, b_end in matcher.get_opcodes():
        old_block = original_lines[a_start:a_end]
        new_block = revised_lines[b_start:b_end]
        if tag == "equal":
            unchanged_lines += len(old_block)
        elif tag == "insert":
            added.extend(new_block)
        elif tag == "delete":
            removed.extend(old_block)
        elif tag == "replace":
            paired = min(len(old_block), len(new_block))
            modified.extend(
                {"before": old_block[index], "after": new_block[index]}
                for index in range(paired)
            )
            removed.extend(old_block[paired:])
            added.extend(new_block[paired:])

    unified = "\n".join(
        difflib.unified_diff(
            original_lines,
            revised_lines,
            fromfile=original_name,
            tofile=revised_name,
            lineterm="",
        )
    )
    return BasicDiffResult(
        added=added,
        removed=removed,
        modified=modified,
        unified_diff=unified,
        similarity=matcher.ratio(),
        unchanged_lines=unchanged_lines,
    )


_BLOCKED_HOSTNAMES = {
    "localhost",
    "localhost.localdomain",
    "metadata.google.internal",
    "metadata.azure.internal",
}


def _is_public_ip(address: str) -> bool:
    try:
        ip = ipaddress.ip_address(address)
        return (
            ip.is_global
            and not ip.is_multicast
            and not ip.is_reserved
            and not ip.is_unspecified
            and not ip.is_loopback
            and not ip.is_link_local
            and not ip.is_private
        )
    except ValueError:
        return False


@lru_cache(maxsize=64)
def _resolve_hostname(hostname: str, port: int) -> tuple[str, ...]:
    addresses = {
        result[4][0]
        for result in socket.getaddrinfo(
            hostname,
            port,
            type=socket.SOCK_STREAM,
        )
    }
    return tuple(sorted(addresses))


def validate_base_url(
    base_url: str,
    *,
    allow_private: bool = False,
    resolve_dns: bool = True,
    require_https: bool = True,
) -> str:
    """验证 OpenAI-compatible Base URL，并阻止 BYOK 请求访问内部网络。"""
    normalized = str(base_url or "").strip().rstrip("/")
    parsed = urlparse(normalized)
    if parsed.scheme not in ("http", "https") or not parsed.netloc:
        raise AIConfigurationError("Base URL 必须是有效的 HTTP(S) 地址。")
    if require_https and parsed.scheme != "https":
        raise AIConfigurationError("BYOK Base URL 必须使用 HTTPS。")
    if parsed.username or parsed.password:
        raise AIConfigurationError("Base URL 不能包含用户名或密码。")
    if parsed.query or parsed.fragment:
        raise AIConfigurationError("Base URL 不能包含查询参数或片段。")

    hostname = (parsed.hostname or "").lower().rstrip(".")
    if not hostname:
        raise AIConfigurationError("Base URL 缺少有效主机名。")
    if not allow_private:
        if (
            hostname in _BLOCKED_HOSTNAMES
            or hostname.endswith((".localhost", ".local", ".internal"))
        ):
            raise AIConfigurationError("Base URL 不能指向本机、内网或元数据服务。")

        try:
            direct_ip = ipaddress.ip_address(hostname)
        except ValueError:
            direct_ip = None
        if direct_ip is not None:
            if not _is_public_ip(str(direct_ip)):
                raise AIConfigurationError("Base URL 不能指向私有、保留或本地 IP。")
        elif resolve_dns:
            try:
                addresses = _resolve_hostname(
                    hostname,
                    parsed.port or (443 if parsed.scheme == "https" else 80),
                )
            except OSError as exc:
                raise AIConfigurationError("Base URL 主机名无法解析，请检查地址。") from exc
            if not addresses or any(not _is_public_ip(address) for address in addresses):
                raise AIConfigurationError("Base URL 解析到了私有、保留或本地 IP。")
    return normalized


def resolve_ai_config(
    server_api_key: str,
    server_base_url: str,
    byok_api_key: str,
    byok_base_url: str,
    model: str,
    server_model: Optional[str] = None,
) -> AIConfig:
    """服务器 Key 始终绑定服务器 URL/Model；仅 BYOK 可使用前端配置。"""
    server_key = str(server_api_key or "").strip()
    if server_key:
        return AIConfig(
            api_key=server_key,
            base_url=validate_base_url(
                server_base_url or DEFAULT_BASE_URL,
                allow_private=True,
                resolve_dns=False,
                require_https=False,
            ),
            model=str(server_model or "deepseek-chat").strip() or "deepseek-chat",
            source="server",
        )

    user_key = str(byok_api_key or "").strip()
    if not user_key:
        raise AIConfigurationError("请先输入 API Key。")
    return AIConfig(
        api_key=user_key,
        base_url=validate_base_url(byok_base_url or DEFAULT_BASE_URL),
        model=str(model or "deepseek-chat").strip() or "deepseek-chat",
        source="byok",
    )


def extract_completion_result(response: Any) -> CompletionResult:
    """兼容 SDK 对象/字典，统一提取内容和 finish_reason。"""
    choices = response.get("choices") if isinstance(response, Mapping) else getattr(response, "choices", None)
    if not choices:
        raise AIResponseError("AI 服务没有返回候选结果。")

    choice = choices[0]
    if isinstance(choice, Mapping):
        message = choice.get("message") or {}
        text = message.get("content", "") if isinstance(message, Mapping) else getattr(message, "content", "")
        finish_reason = choice.get("finish_reason") or ""
    else:
        message = getattr(choice, "message", None)
        text = getattr(message, "content", "") if message is not None else ""
        finish_reason = getattr(choice, "finish_reason", "") or ""

    normalized = str(text or "").strip()
    if not normalized:
        raise AIResponseError("AI 服务返回了空内容，请稍后重试或更换模型。")
    return CompletionResult(normalized, str(finish_reason))


def complete_with_continuations(
    request: Callable[[Sequence[Mapping[str, Any]]], CompletionResult],
    messages: Sequence[Mapping[str, Any]],
    max_continuations: int = MAX_TRANSFORM_CONTINUATIONS,
) -> ContinuationResult:
    """对 length 响应有限续写，避免翻译/改写静默丢失后半部分。"""
    if max_continuations < 0:
        raise ValueError("max_continuations 不能小于 0。")
    working_messages = [dict(message) for message in messages]
    parts: List[str] = []
    for attempt in range(max_continuations + 1):
        result = request(working_messages)
        parts.append(result.text)
        if not result.truncated:
            return ContinuationResult("\n".join(parts), False, attempt)
        if attempt >= max_continuations:
            return ContinuationResult("\n".join(parts), True, attempt)
        working_messages.extend(
            [
                {"role": "assistant", "content": result.text},
                {
                    "role": "user",
                    "content": (
                        "上次输出因长度限制中断。请从中断处直接继续，不要重复已有内容，"
                        "不要新增总结，保持原顺序与格式。"
                    ),
                },
            ]
        )
    return ContinuationResult("\n".join(parts), True, max_continuations)


def estimate_token_usage(messages: Sequence[Mapping[str, Any]], output: str) -> TokenUsage:
    prompt_text = "\n".join(str(message.get("content", "")) for message in messages)
    prompt_tokens = max(1, len(prompt_text) // 3) if prompt_text else 0
    completion_tokens = max(1, len(output) // 3) if output else 0
    return TokenUsage(
        prompt_tokens=prompt_tokens,
        completion_tokens=completion_tokens,
        total_tokens=prompt_tokens + completion_tokens,
        estimated=True,
    )


def extract_token_usage(response: Any) -> Optional[TokenUsage]:
    """兼容 OpenAI/DeepSeek SDK 对象和字典形式的 response.usage。"""
    usage = getattr(response, "usage", None)
    if usage is None and isinstance(response, Mapping):
        usage = response.get("usage")
    if not usage:
        return None

    def _value(name: str) -> Optional[int]:
        raw = usage.get(name) if isinstance(usage, Mapping) else getattr(usage, name, None)
        try:
            return int(raw) if raw is not None else None
        except (TypeError, ValueError):
            return None

    prompt_tokens = _value("prompt_tokens")
    completion_tokens = _value("completion_tokens")
    total_tokens = _value("total_tokens")
    if prompt_tokens is None and completion_tokens is None and total_tokens is None:
        return None
    prompt_tokens = prompt_tokens or 0
    completion_tokens = completion_tokens or 0
    total_tokens = total_tokens if total_tokens is not None else prompt_tokens + completion_tokens
    return TokenUsage(prompt_tokens, completion_tokens, total_tokens, estimated=False)


def detect_network_environment() -> str:
    try:
        with socket.create_connection(("www.google.com", 443), timeout=3):
            return "global"
    except OSError as exc:
        logger.info("global connectivity check unavailable [%s]", type(exc).__name__)
        return "china"


def _search_tavily(query: str, api_key: str) -> List[Dict[str, str]]:
    import httpx

    response = httpx.post(
        "https://api.tavily.com/search",
        json={"api_key": api_key, "query": query, "max_results": 5},
        timeout=15,
    )
    response.raise_for_status()
    payload = response.json()
    return [
        {
            "title": str(item.get("title", "")),
            "url": str(item.get("url", "")),
            "content": str(item.get("content", ""))[:400],
            "engine": "Tavily",
        }
        for item in payload.get("results", [])
        if isinstance(item, Mapping)
    ]


def _search_duckduckgo(query: str) -> List[Dict[str, str]]:
    from ddgs import DDGS

    results = DDGS().text(query, max_results=5)
    return [
        {
            "title": str(item.get("title", "")),
            "url": str(item.get("href", item.get("url", ""))),
            "content": str(item.get("body", item.get("content", "")))[:400],
            "engine": "DuckDuckGo",
        }
        for item in results
        if isinstance(item, Mapping)
    ]


def _search_bocha(query: str, api_key: str) -> List[Dict[str, str]]:
    import httpx

    response = httpx.post(
        "https://api.bochaai.com/v1/web-search",
        headers={"Authorization": f"Bearer {api_key}"},
        json={"query": query, "count": 5, "summary": True},
        timeout=20,
    )
    response.raise_for_status()
    pages = response.json().get("data", {}).get("webPages", {}).get("value", [])
    return [
        {
            "title": str(page.get("name", "")),
            "url": str(page.get("url", "")),
            "content": str(page.get("summary", page.get("snippet", "")))[:400],
            "engine": "博查",
        }
        for page in pages
        if isinstance(page, Mapping)
    ]


def _search_baidu(query: str) -> List[Dict[str, str]]:
    url = f"https://www.baidu.com/s?wd={urllib.parse.quote(query)}&rn=8"
    request = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
    with urllib.request.urlopen(request, timeout=10) as response:
        response_text = response.read().decode("utf-8", errors="replace")

    results: List[Dict[str, str]] = []
    pattern = re.compile(
        r'<h3[^>]*>.*?<a[^>]+href="([^"]+)"[^>]*>(.*?)</a>', re.DOTALL
    )
    for match in pattern.finditer(response_text):
        title = re.sub(r"<[^>]+>", "", match.group(2))
        title = html.unescape(title).strip()
        href = match.group(1)
        if title and len(title) > 3:
            if href.startswith("/link"):
                href = "https://www.baidu.com" + href
            results.append(
                {"title": title, "url": href, "content": "", "engine": "百度"}
            )
        if len(results) >= 5:
            break
    return results


def search_web(
    query: str,
    network_mode: str = "auto",
    tavily_key: str = "",
    bocha_key: str = "",
) -> SearchOutcome:
    """按网络环境尝试搜索服务，并显式返回回退或失败信息。"""
    normalized_query = str(query or "").strip()
    if not normalized_query:
        return SearchOutcome([], error="请输入搜索关键词。")

    environment = (
        detect_network_environment() if network_mode == "auto" else network_mode
    )
    attempts = []
    notices: List[str] = []

    if environment == "global":
        if tavily_key:
            attempts.append(("Tavily", lambda: _search_tavily(normalized_query, tavily_key)))
        attempts.append(("DuckDuckGo", lambda: _search_duckduckgo(normalized_query)))
    else:
        if bocha_key:
            attempts.append(("博查", lambda: _search_bocha(normalized_query, bocha_key)))
        attempts.append(("百度", lambda: _search_baidu(normalized_query)))

    failed_engines: List[str] = []
    for engine, searcher in attempts:
        try:
            results = searcher()
            if results:
                if failed_engines:
                    notices.append(
                        f"{', '.join(failed_engines)} 暂不可用，已自动切换到 {engine}。"
                    )
                return SearchOutcome(results, notices=notices)
            failed_engines.append(engine)
            logger.warning("search provider returned no results: %s", engine)
        except Exception as exc:
            failed_engines.append(engine)
            log_failure(f"search provider {engine}", exc)

    return SearchOutcome(
        [],
        notices=notices,
        error="联网搜索暂时不可用，请稍后重试或切换网络模式。",
    )
